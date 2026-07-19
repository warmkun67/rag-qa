import os
import pandas as pd
from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    Docx2txtLoader,
    UnstructuredExcelLoader,
    UnstructuredMarkdownLoader
)
from langchain_text_splitters import RecursiveCharacterTextSplitter
import config

def load_all_docs():
    docs = []
    if not os.path.exists(config.DOCS_DIR):
        print(f"警告：文档目录 {config.DOCS_DIR} 不存在")
        return docs

    print(f"开始加载 {config.DOCS_DIR} 目录下的文档...")
    for filename in os.listdir(config.DOCS_DIR):
        path = os.path.join(config.DOCS_DIR, filename)
        print(f"正在处理文件：{filename}")
        try:
            # PDF 加载
            if filename.endswith(".pdf"):
                loader = PyPDFLoader(path)
                pdf_docs = loader.load()
                print(f"✅ 成功加载PDF：{filename}，共 {len(pdf_docs)} 页")
                docs.extend(pdf_docs)

            # TXT 加载
            elif filename.endswith(".txt"):
                loader = TextLoader(path, encoding="utf-8")
                txt_docs = loader.load()
                print(f"✅ 成功加载TXT：{filename}")
                docs.extend(txt_docs)

            # Word 加载（增强兼容）
            elif filename.endswith(".docx"):
                # 优先用 Docx2txtLoader，失败则用 python-docx 兜底
                try:
                    loader = Docx2txtLoader(path)
                    docx_docs = loader.load()
                    print(f"✅ 成功加载DOCX：{filename}")
                    docs.extend(docx_docs)
                except ImportError:
                    # 兜底方案：用 python-docx 直接读取
                    from docx import Document
                    doc = Document(path)
                    full_text = "\n".join([para.text for para in doc.paragraphs])
                    from langchain_core.documents import Document
                    docx_docs = [Document(page_content=full_text)]
                    print(f"✅ 成功加载DOCX：{filename}")
                    docs.extend(docx_docs)

            # Markdown 加载
            elif filename.endswith(".md"):
                loader = UnstructuredMarkdownLoader(path)
                md_docs = loader.load()
                print(f"✅ 成功加载MD：{filename}")
                docs.extend(md_docs)

            # Excel 加载
            elif filename.endswith((".xlsx", ".xls")):
                loader = UnstructuredExcelLoader(path, mode="elements")
                excel_docs = loader.load()
                print(f"✅ 成功加载Excel：{filename}")
                docs.extend(excel_docs)

            # CSV 加载
            elif filename.endswith(".csv"):
                df = pd.read_csv(path)
                text = df.to_string(index=False)
                from langchain_core.documents import Document
                csv_doc = Document(page_content=text)
                print(f"✅ 成功加载CSV：{filename}")
                docs.append(csv_doc)

            else:
                print(f"⚠️ 跳过不支持的文件格式：{filename}")

        except Exception as e:
            print(f"❌ 加载文件 {filename} 失败，错误信息：{str(e)}")
            continue

    print(f"所有文档加载完成，共 {len(docs)} 个文档块，开始分块...")
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1024,
        chunk_overlap=128,
        separators=["\n\n", "\n", "。", "！", "？", " "]
    )

    splits = splitter.split_documents(docs)
    print(f"分块完成，共 {len(splits)} 个文本块")
    return splits